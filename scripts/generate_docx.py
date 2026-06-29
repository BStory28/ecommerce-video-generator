import json, os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_json(path):
    with open(path, encoding="utf-8") as f:
        return json.load(f)

def gen_docx(storyboard_path, output_dir):
    sb = load_json(storyboard_path)
    shots = sb.get("镜头脚本", [])

    product_name = list(sb.keys())[0].replace("产品参考图", "").strip()
    ref_img_line = sb.get(list(sb.keys())[0], "")
    people_desc = sb.get("人物形象特征参考", "")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)

    # ── 标题 ──
    h = doc.add_heading(f"【{product_name}】泰国市场 UGC种草 分镜脚本", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ═══════════════════════════════════════════
    # 一、基础信息
    # ═══════════════════════════════════════════
    doc.add_heading("一、基础信息", level=1)

    # 1. 产品参考图
    p = doc.add_paragraph()
    run = p.add_run("1. 产品参考图")
    run.bold = True
    p.add_run(f"：{ref_img_line}")

    # 2. 目标人群
    p = doc.add_paragraph()
    run = p.add_run("2. 目标人群")
    run.bold = True
    p.add_run(f"：{people_desc}")

    # 3. 产品核心属性
    pf = sb.get("产品功能属性", "")
    pp = sb.get("用户痛点", "")
    core_desc = (
        f"主打效果对比，直击吹宠时毛发四处乱飞、难清理的用户痛点；"
        f"网状围笼结构，可隔绝散落毛发，操作简单，宠物不易应激，透气不闷热，快速吹干宠物毛发。"
    )
    p = doc.add_paragraph()
    run = p.add_run("3. 产品核心属性")
    run.bold = True
    p.add_run(f"：{core_desc}")

    # ═══════════════════════════════════════════
    # 二、分镜脚本
    # ═══════════════════════════════════════════
    total_dur = sum(int(s.get("单镜时长","0s").replace("s","")) for s in shots)
    doc.add_heading(f"二、分镜脚本（总时长{total_dur}s）", level=1)

    for s in shots:
        shot_num = s.get("镜号", "?")
        dur = s.get("单镜时长", "?")
        scene = s.get("景别", "")
        angle = s.get("拍摄角度", "")
        camera = s.get("运镜", "")
        content = s.get("画面内容", "")
        dialogue = s.get("本地台词", "")
        subtitle = s.get("本地字幕", "")
        audio = s.get("音效_BGM", "")
        lighting = s.get("光影色调", "")
        realism = s.get("真实感约束", "")

        # 镜号标题
        title_parts = [f"镜号{shot_num}｜{dur}"]
        if scene:
            title_parts.append(scene)
        if angle:
            title_parts.append(angle)
        if camera:
            title_parts.append(camera)
        doc.add_heading("｜".join(title_parts), level=2)

        # 画面
        p = doc.add_paragraph()
        run = p.add_run("画面")
        run.bold = True
        p.add_run(f"：{content}")

        # 台词
        p = doc.add_paragraph()
        run = p.add_run("泰语台词")
        run.bold = True
        p.add_run(f"：{dialogue}")

        # 字幕
        p = doc.add_paragraph()
        run = p.add_run("泰语字幕")
        run.bold = True
        p.add_run(f"：{subtitle}")

        # 音效BGM
        p = doc.add_paragraph()
        run = p.add_run("音效BGM")
        run.bold = True
        p.add_run(f"：{audio}")

        # 光影色调
        p = doc.add_paragraph()
        run = p.add_run("光影色调")
        run.bold = True
        p.add_run(f"：{lighting}")

        # 真实约束
        p = doc.add_paragraph()
        run = p.add_run("真实约束")
        run.bold = True
        p.add_run(f"：{realism}")

    # ── 保存 ──
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, f"{product_name}_thailand_UGC种草.docx")
    doc.save(out_path)
    print(out_path)

if __name__ == "__main__":
    gen_docx(sys.argv[1], sys.argv[2])
