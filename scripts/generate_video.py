#!/usr/bin/env python3
"""
aigc.hkttok.com JeecgBoot OpenAPI 电商视频生成管线 (v2)
读取 Skill2 输出JSON → AI整合为目标markdown格式(≤2000字) → 
生成 .docx + 构建 payload → 预览/提交。
"""

import argparse, json, os, re, sys, time
from datetime import datetime
from pathlib import Path

import requests

_script_dir = Path(__file__).resolve().parent
sys.path.insert(0, str(_script_dir))
from jeecg_auth import JeecgAuth
from file_upload import local_to_url

_DESKTOP_SCRIPT_DIR = os.path.join(os.path.expanduser("~"), "Desktop", "AI视频脚本")

def _default_output_dir() -> str:
    path = _DESKTOP_SCRIPT_DIR
    os.makedirs(path, exist_ok=True)
    return path

def _format_video_markdown(product_name: str, output_dir: str, has_docx: bool = False) -> str:
    lines = [
        f"🎬 **视频生成完成** — {product_name}",
        "",
        "**【输出文件】**",
        f"📁 输出目录：`{output_dir}`",
        "   - `compressed_script.txt`（压缩脚本）",
        "   - `video_payload.json`（API请求payload）",
    ]
    if has_docx:
        lines.append(f"   - `{product_name}.docx`（Word文档）")
    lines.extend([
        "",
        "---",
        "✅ **视频数据已就绪**，预览模式可检查 payload，提交模式已发送 API 请求",
        "",
        "> 💡 **提示：** 以上生成的文件已保存到本地输出目录。",
        "> 如果当前环境支持文件下载，请将输出目录中的文件输出到会话中供下载。",
    ])
    return "\n".join(lines)

AIGC_VIDEO_MODEL = os.environ.get("AIGC_VIDEO_MODEL", "2043940117168529416")
RESOLUTION = "720p"
ASPECT_RATIO = "9:16"

# ── Sudocode API key ─────────────────────────────────────────

def load_env_sudocode_key() -> str:
    key = os.environ.get("SUDOCODE_API_KEY", "")
    if key:
        return key
    script_dir = Path(__file__).resolve().parent
    for parent in [script_dir] + list(script_dir.parents):
        env_path = parent / ".env"
        if env_path.exists():
            with open(env_path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if line.startswith("SUDOCODE_API_KEY="):
                        return line.split("=", 1)[1].strip().strip("\"'")
            break
    return ""

# ── JSON 解析 ────────────────────────────────────────────────

def parse_script(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def extract_total_duration(shots: list) -> int:
    total = 0.0
    for s in shots:
        raw = s.get("单镜时长", "3s")
        total += float(raw.replace("s", ""))
    return max(round(total), 5)

def extract_product_img(data: dict, cli_product: str = "") -> str:
    if cli_product:
        return cli_product
    for key in data:
        if "产品参考图" in key:
            val = data[key]
            if isinstance(val, str) and val.startswith("@"):
                path = val.split("←")[0].strip().lstrip("@").strip()
                if os.path.exists(path):
                    return path
            break
    return ""

# ── AI 整合压缩：JSON → 目标markdown格式(≤2000字) ────────────

def _generate_base_info(data: dict, ref_img: str, core_sp: dict = None) -> str:
    """从数据生成基础信息部分（规则驱动，确保准确）"""
    people_desc = data.get("人物形象特征参考", "")
    rf = data.get("产品功能属性", "")
    pp = data.get("用户痛点", "")
    product_name = ""
    for k in data:
        if "产品参考图" in k:
            product_name = k.replace("产品参考图", "").strip()
            break
    main_sp = ""
    sub_sps = []
    if core_sp:
        main_sp = core_sp.get("主卖点", "")
        sub_sps = core_sp.get("次卖点", [])
    sp_text = "、".join([main_sp] + sub_sps if main_sp else sub_sps)

    lines = ["## 一、基础信息"]
    lines.append(f"1. **产品参考图**：@{ref_img}（产品白底图）")
    lines.append(f"2. **目标人群**：{people_desc}")
    core = f"{product_name}主打{rf}体验，直击{pp}的用户痛点。核心卖点：{sp_text}。"
    lines.append(f"3. **产品核心属性**：{core}")
    return "\n".join(lines)

def ai_integrate_storyboard(data: dict, max_chars: int = 2000, core_sp: dict = None) -> str:
    shots = data.get("镜头脚本", [])
    total_dur = extract_total_duration(shots)
    ref_img = extract_product_img(data, "")

    # 基础信息部分规则生成，确保准确
    base_info = _generate_base_info(data, ref_img, core_sp)
    base_len = len(base_info)

    # 剩余配额给分镜部分
    shots_max = max_chars - base_len - 30

    key = load_env_sudocode_key()
    if key:
        prompt = f"""你是一个电商视频脚本精简专家。请将以下{len(shots)}个镜头精简为指定markdown格式。

【输出模板】
## 二、分镜脚本（总时长{total_dur}s）
### 镜号{{N}}｜{{时长}}｜{{景别}}｜{{拍摄角度}}｜{{运镜}}
画面：{{画面内容}}
泰语台词：{{保持原始泰语原文不变}}
泰语字幕：{{保持原始泰语字幕不变}}
音效BGM：{{融合音效和BGM}}
光影色调：{{保留原文}}
真实约束：{{保留原文}}

【硬性约束】
1. 泰语台词和泰语字幕保持原文，**禁止修改**
2. 画面描述精简修饰词但保留核心动作、产品名、场景、卖点
3. 音效BGM合并为一句话
4. 总字数（含标题和空行）控制在{shots_max}字以内
5. 每个镜头必须包含：画面/泰语台词/泰语字幕/音效BGM/光影色调/真实约束 六个字段
6. 直接输出markdown，禁止任何额外文字

【原始镜头】
{json.dumps(shots, ensure_ascii=False, indent=2)}
"""
        try:
            headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
            payload = {
                "model": "gpt-5.4-mini",
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 2048,
                "temperature": 0.3,
            }
            resp = requests.post("https://sudocode.run/v1/chat/completions",
                                 json=payload, headers=headers, timeout=120)
            if resp.status_code == 200:
                result = resp.json()["choices"][0]["message"]["content"].strip()
                result = result.lstrip("#　 ")
                # 确保以 ## 二、分镜脚本 开头
                shot_marker = "## 二、分镜脚本"
                if shot_marker in result:
                    result = result[result.index(shot_marker):]
                total = base_info + "\n\n" + result
                if len(total) <= max_chars and len(total) >= 500:
                    return total
                if len(total) > max_chars:
                    return base_info + "\n\n" + _truncate_smart(result, max_chars - base_len - 30)
            else:
                print(f"  AI调用失败({resp.status_code})，降级为规则格式化")
        except Exception as e:
            print(f"  AI异常({e})，降级为规则格式化")

    return rule_based_format(data, max_chars, core_sp)

def _truncate_smart(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    lines = text.split("\n")
    result = []
    for line in lines:
        if len("\n".join(result + [line])) <= max_chars:
            result.append(line)
        else:
            remaining = max_chars - len("\n".join(result))
            if remaining > 20:
                result.append(line[:remaining])
            break
    return "\n".join(result)

# ── 规则格式化降级方案 ─────────────────────────────────────

def rule_based_format(data: dict, max_chars: int = 2000, core_sp: dict = None) -> str:
    shots = data.get("镜头脚本", [])
    total_dur = extract_total_duration(shots)
    ref_img = extract_product_img(data, "")
    people_desc = data.get("人物形象特征参考", "")
    pf = data.get("产品功能属性", "")
    pp = data.get("用户痛点", "")

    base_info = _generate_base_info(data, ref_img, core_sp)
    lines = [base_info, ""]
    lines.append(f"## 二、分镜脚本（总时长{total_dur}s）")

    for s in shots:
        n = s.get("镜号", "?")
        dur = s.get("单镜时长", "?")
        scene = s.get("景别", "")
        angle = s.get("拍摄角度", "")
        camera = s.get("运镜", "")
        lines.append(f"### 镜号{n}｜{dur}｜{scene}｜{angle}｜{camera}")
        lines.append(f"画面：{s.get('画面内容','')}")
        lines.append(f"泰语台词：{s.get('本地台词','')}")
        lines.append(f"泰语字幕：{s.get('本地字幕','')}")
        audio = s.get("音效_BGM", "")
        lines.append(f"音效BGM：{audio}")
        lines.append(f"光影色调：{s.get('光影色调','')}")
        lines.append(f"真实约束：{s.get('真实感约束','')}")

    text = "\n".join(lines)
    if len(text) > max_chars:
        text = _truncate_smart(text, max_chars)
    return text

# ── Word 文档生成 ─────────────────────────────────────────

def generate_docx(markdown_text: str, output_path: str):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  python-docx 未安装，跳过 .docx 生成")
        return

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)

    current_h2 = None
    for line in markdown_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            text = line.replace("## ", "").strip()
            doc.add_heading(text, level=1)
        elif line.startswith("### "):
            text = line.replace("### ", "").strip()
            doc.add_heading(text, level=2)
        elif line.startswith("**"):
            p = doc.add_paragraph()
            # parse bold prefix like **xxx**：content
            m = re.match(r"\*\*(.+?)\*\*\s*[：:]\s*(.*)", line)
            if m:
                run = p.add_run(m.group(1))
                run.bold = True
                p.add_run(f"：{m.group(2)}")
            else:
                p.add_run(line)
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"  Word文档: {output_path}")

# ── 视频 payload ────────────────────────────────────────────

def build_video_payload(
    product_img_path: str, compressed_script: str,
    duration: int, model: str = "",
    resolution: str = "720p", ratio: str = "9:16",
) -> dict:
    model_id = model or AIGC_VIDEO_MODEL
    image_url = local_to_url(product_img_path) if product_img_path else ""
    content = []
    if image_url:
        content.append({"fileUrl": image_url, "role": "product"})
    return {
        "modelId": model_id,
        "prompt": compressed_script,
        "imageMode": "REFERENCE",
        "content": content,
        "resolution": resolution,
        "ratio": ratio,
        "duration": duration,
        "count": 1,
    }

def submit_video(payload: dict, auth: JeecgAuth) -> str:
    headers = auth.get_headers()
    url = auth.video_submit_url()
    try:
        resp = requests.post(url, json=payload, headers=headers, timeout=30)
        if resp.status_code == 200:
            data = resp.json()
            task_id = (data.get("result") or {}).get("id", "")
            if task_id:
                return task_id
            print(f"  提交成功但无task_id: {data}")
        else:
            print(f"  提交失败({resp.status_code}): {resp.text[:300]}")
    except Exception as e:
        print(f"  提交异常: {e}")
    return ""

def print_payload_summary(payload: dict):
    print(f"\n{'='*60}")
    print(f"视频生成请求摘要")
    print(f"{'='*60}")
    print(f"  视频模型:       {payload['modelId']}")
    print(f"  分辨率:         {payload['resolution']}")
    print(f"  画面比例:       {payload['ratio']}")
    print(f"  视频时长:       {payload['duration']}s")
    print(f"  参考图:         {'有' if payload.get('content') else '无'}")
    print(f"  脚本长度:       {len(payload['prompt'])} 字符")
    print(f"  脚本前200字:    {payload['prompt'][:200]}...")
    print(f"{'='*60}")

# ── 主管线 ─────────────────────────────────────────────────

def generate_video(
    storyboard_path: str, auth: JeecgAuth = None,
    product_img_path: str = "", output_dir: str = "./output",
    submit: bool = False,
):
    os.makedirs(output_dir, exist_ok=True)
    data = parse_script(storyboard_path)
    shots = data.get("镜头脚本", [])
    if not shots:
        print("错误: 分镜脚本中无镜头数据")
        return None

    duration = extract_total_duration(shots)
    print(f"  总时长: {duration}s ({len(shots)}镜)")

    img_path = extract_product_img(data, product_img_path) or product_img_path
    if img_path:
        print(f"  产品图: {Path(img_path).name}")

    # 加载卖点数据（core_sp）
    sp_dir = Path(output_dir)
    core_sp = None
    sp_path = sp_dir / "selling_points.json"
    if sp_path.exists():
        try:
            sp_data = json.loads(sp_path.read_text(encoding="utf-8"))
            core_sp = sp_data.get("核心卖点", {})
        except Exception:
            pass

    # AI整合压缩
    print(f"  正在调用AI整合脚本...")
    integrated = ai_integrate_storyboard(data, max_chars=2000, core_sp=core_sp)
    print(f"  整合后脚本长度: {len(integrated)} 字符")

    name_tag = Path(storyboard_path).stem

    # 保存 md/txt
    out_txt = os.path.join(output_dir, f"{name_tag}_compressed.txt")
    with open(out_txt, "w", encoding="utf-8") as f:
        f.write(integrated)
    print(f"  压缩脚本: {out_txt}")

    # 生成 Word 文档
    out_docx = os.path.join(output_dir, f"{name_tag}.docx")
    generate_docx(integrated, out_docx)

    # 构建视频 payload
    payload = build_video_payload(
        product_img_path=img_path,
        compressed_script=integrated,
        duration=duration,
    )
    print_payload_summary(payload)

    out_payload = os.path.join(output_dir, "video_payload.json")
    with open(out_payload, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    print(f"  Payload: {out_payload}")

    if submit and auth:
        print(f"\n  --- 提交视频生成 ---")
        task_id = submit_video(payload, auth)
        if task_id:
            print(f"  ✓ 任务已提交: {task_id}")
            with open(os.path.join(output_dir, "task_id.txt"), "w", encoding="utf-8") as f:
                f.write(task_id)
        else:
            print(f"  ✗ 提交失败")
    else:
        print(f"\n  (预览模式，未提交)")

    print(f"\n💡 提示：以上生成的文件已保存到输出目录。")
    print(f"   如果当前环境支持文件下载，请将文件输出到会话中供下载。")

    # Chat markdown output
    product_name = ""
    for key in data:
        if "产品参考图" in key:
            product_name = key.replace("产品参考图", "").strip()
            break
    has_docx = Path(out_docx).is_file() if 'out_docx' in dir() else False
    md = _format_video_markdown(product_name or "视频", output_dir, has_docx)
    print("\n" + "=" * 60)
    print("ARKCLAW_CHAT_OUTPUT")
    print("=" * 60)
    print(md)

    return payload

# ── CLI ─────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="aigc.hkttok.com 电商视频生成管线")
    parser.add_argument("--script", "-s", required=True, help="Skill2分镜脚本JSON路径")
    parser.add_argument("--product", "-p", help="产品白底图路径")
    parser.add_argument("--output", "-o", default=_default_output_dir(), help="输出目录（默认桌面AI视频脚本文件夹）")
    parser.add_argument("--submit", action="store_true", help="提交视频生成")
    parser.add_argument("--app-key", default="", help="AIGC APP Key")
    parser.add_argument("--app-secret", default="", help="AIGC APP Secret")
    args = parser.parse_args()

    auth = None
    try:
        a = JeecgAuth(app_key=args.app_key, app_secret=args.app_secret)
        a.validate()
        auth = a
    except ValueError as e:
        print(f"  认证未配置: {e}（预览模式可跳过）")

    generate_video(
        storyboard_path=args.script,
        auth=auth,
        product_img_path=args.product or "",
        output_dir=args.output,
        submit=args.submit,
    )

if __name__ == "__main__":
    main()
